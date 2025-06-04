from alpha_vantage.timeseries import TimeSeries
from alpha_vantage.fundamentaldata import FundamentalData
import pandas as pd
import numpy as np
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import plotly.graph_objects as go
from typing import Dict, List, Tuple
import os
from dotenv import load_dotenv

class StockOnePager:
    def __init__(self, ticker: str):
        """
        Initialize the Stock One-Pager Generator for a given ticker.
        
        Args:
            ticker (str): Stock ticker symbol
        """
        load_dotenv()
        self.ticker = ticker.upper()
        self.alpha_vantage_key = os.getenv('ALPHA_VANTAGE_API_KEY')
        self.ts = TimeSeries(key=self.alpha_vantage_key, output_format='pandas')
        self.fd = FundamentalData(key=self.alpha_vantage_key)
        # Fetch company info
        self.info = self._get_company_info()
        # Fetch historical data
        self.historical_data = self._get_historical_data()
        
    def _get_company_info(self) -> Dict:
        try:
            data, _ = self.fd.get_company_overview(self.ticker)
            return data
        except Exception as e:
            print(f"Error fetching company info: {e}")
            return {}

    def _get_historical_data(self) -> pd.DataFrame:
        try:
            data, _ = self.ts.get_daily(symbol=self.ticker, outputsize='full')
            data = data.rename(columns={
                '1. open': 'Open',
                '2. high': 'High',
                '3. low': 'Low',
                '4. close': 'Close',
                '5. volume': 'Volume'
            })
            cutoff = datetime.now() - timedelta(days=5*365)
            data = data[data.index >= cutoff]
            return data
        except Exception as e:
            print(f"Error fetching historical data: {e}")
            return pd.DataFrame()

    def calculate_valuation_metrics(self) -> Dict:
        """
        Calculate key valuation metrics for the stock.
        
        Returns:
            Dict: Dictionary containing valuation metrics
        """
        try:
            # Basic metrics
            current_price = float(self.historical_data['Close'].iloc[-1]) if not self.historical_data.empty else 0
            market_cap = float(self.info.get('MarketCapitalization', 0))
            pe_ratio = float(self.info.get('PERatio', 0))
            forward_pe = float(self.info.get('ForwardPE', 0)) if self.info.get('ForwardPE') else 0
            eps = float(self.info.get('EPS', 0)) if self.info.get('EPS') else 0
            
            # Growth metrics
            revenue_growth = float(self.info.get('QuarterlyRevenueGrowthYOY', 0))
            earnings_growth = float(self.info.get('QuarterlyEarningsGrowthYOY', 0))
            
            # Financial health metrics
            debt_to_equity = float(self.info.get('DebtToEquity', 0))
            current_ratio = float(self.info.get('CurrentRatio', 0))
            quick_ratio = float(self.info.get('QuickRatio', 0))
            
            # Dividend metrics
            dividend_yield = float(self.info.get('DividendYield', 0))
            
            return {
                'current_price': current_price,
                'market_cap': market_cap,
                'pe_ratio': pe_ratio,
                'forward_pe': forward_pe,
                'eps': eps,
                'revenue_growth': revenue_growth,
                'earnings_growth': earnings_growth,
                'debt_to_equity': debt_to_equity,
                'current_ratio': current_ratio,
                'quick_ratio': quick_ratio,
                'dividend_yield': dividend_yield
            }
        except Exception as e:
            print(f"Error calculating valuation metrics: {e}")
            return {}
    
    def generate_growth_one_pager(self) -> Document:
        """
        Generate a Growth-focused one-pager.
        
        Returns:
            Document: Word document containing the growth one-pager
        """
        doc = Document()
        metrics = self.calculate_valuation_metrics()
        
        # Add title
        title = doc.add_heading(f'{self.ticker} - Growth Analysis', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add company overview
        doc.add_heading('Company Overview', level=1)
        doc.add_paragraph(self.info.get('longBusinessSummary', 'No business summary available.'))
        
        # Add growth metrics
        doc.add_heading('Growth Metrics', level=1)
        growth_metrics = [
            ('Revenue Growth', f"{metrics['revenue_growth']*100:.2f}%"),
            ('Earnings Growth', f"{metrics['earnings_growth']*100:.2f}%"),
            ('Forward P/E', f"{metrics['forward_pe']:.2f}"),
            ('Market Cap', f"${metrics['market_cap']/1e9:.2f}B")
        ]
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Value'
        
        for metric, value in growth_metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
            
        return doc
    
    def generate_value_one_pager(self) -> Document:
        """
        Generate a Value-focused one-pager.
        
        Returns:
            Document: Word document containing the value one-pager
        """
        doc = Document()
        metrics = self.calculate_valuation_metrics()
        
        # Add title
        title = doc.add_heading(f'{self.ticker} - Value Analysis', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add company overview
        doc.add_heading('Company Overview', level=1)
        doc.add_paragraph(self.info.get('longBusinessSummary', 'No business summary available.'))
        
        # Add value metrics
        doc.add_heading('Value Metrics', level=1)
        value_metrics = [
            ('Current P/E', f"{metrics['pe_ratio']:.2f}"),
            ('Dividend Yield', f"{metrics['dividend_yield']*100:.2f}%"),
            ('Current Ratio', f"{metrics['current_ratio']:.2f}"),
            ('Debt to Equity', f"{metrics['debt_to_equity']:.2f}")
        ]
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Value'
        
        for metric, value in value_metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
            
        return doc
    
    def generate_core_one_pager(self) -> Document:
        """
        Generate a Core-focused one-pager.
        
        Returns:
            Document: Word document containing the core one-pager
        """
        doc = Document()
        metrics = self.calculate_valuation_metrics()
        
        # Add title
        title = doc.add_heading(f'{self.ticker} - Core Analysis', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add company overview
        doc.add_heading('Company Overview', level=1)
        doc.add_paragraph(self.info.get('longBusinessSummary', 'No business summary available.'))
        
        # Add core metrics
        doc.add_heading('Core Metrics', level=1)
        core_metrics = [
            ('Current Price', f"${metrics['current_price']:.2f}"),
            ('Market Cap', f"${metrics['market_cap']/1e9:.2f}B"),
            ('P/E Ratio', f"{metrics['pe_ratio']:.2f}"),
            ('EPS', f"${metrics['eps']:.2f}"),
            ('Quick Ratio', f"{metrics['quick_ratio']:.2f}")
        ]
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Value'
        
        for metric, value in core_metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
            
        return doc
    
    def save_one_pager(self, doc: Document, style: str):
        """
        Save the generated one-pager to a file.
        
        Args:
            doc (Document): The document to save
            style (str): The style of the one-pager (growth, value, or core)
        """
        filename = f"{self.ticker}_{style}_one_pager_{datetime.now().strftime('%Y%m%d')}.docx"
        doc.save(filename)
        print(f"One-pager saved as {filename}")

def generate_all_one_pagers(ticker: str):
    """
    Generate all types of one-pagers for a given ticker.
    
    Args:
        ticker (str): Stock ticker symbol
    """
    generator = StockOnePager(ticker)
    
    # Generate and save all types of one-pagers
    growth_doc = generator.generate_growth_one_pager()
    value_doc = generator.generate_value_one_pager()
    core_doc = generator.generate_core_one_pager()
    
    generator.save_one_pager(growth_doc, 'growth')
    generator.save_one_pager(value_doc, 'value')
    generator.save_one_pager(core_doc, 'core')

if __name__ == "__main__":
    # Example usage
    ticker = input("Enter stock ticker: ")
    generate_all_one_pagers(ticker) 